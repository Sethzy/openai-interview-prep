from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "wiki/domains/OpenAI-SDR-screen-share-safe-prep-doc.md"
OUTPUT = ROOT / "docs/share/OpenAI-SDR-screen-share-safe-prep-google-docs.docx"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_table_width(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def next_numbering_ids(doc: Document):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc: Document) -> int:
    abstract_id, num_id = next_numbering_ids(doc)
    numbering = doc.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(tabs)
    p_pr.append(ind)
    for node in (start, fmt, text, jc, p_pr):
        lvl.append(node)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num_pr(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("[[", "").replace("]]", "")
    text = text.replace("<br>", " ")
    return text.strip()


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [clean_inline(c) for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r"[-:\s]+", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_para(doc: Document, text: str, style: str | None = None, italic: bool = False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if italic:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(clean_inline(text))
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.italic = italic
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(clean_inline(text))
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_number(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph()
    apply_num_pr(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(clean_inline(text))
    run.font.name = "Arial"
    run.font.size = Pt(11)


def parse_numbered_list(lines: list[str], start: int):
    items = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not re.match(r"^\d+\. ", line):
            break
        items.append(re.sub(r"^\d+\. ", "", line))
        i += 1
    return items, i


def add_heading(doc: Document, text: str, level: int):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    run = p.add_run(clean_inline(text))
    run.font.name = "Arial"
    run.font.bold = False
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(67, 67, 67)


def widths_for(cols: int):
    if cols <= 1:
        return [6.5]
    if cols == 2:
        return [2.0, 4.5]
    if cols == 3:
        return [1.45, 2.55, 2.5]
    if cols == 4:
        return [1.25, 1.75, 1.8, 1.7]
    return [6.5 / cols] * cols


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_borders(table)
    widths = widths_for(cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(clean_inline(text))
            run.font.name = "Arial"
            run.font.size = Pt(10 if cols >= 4 else 10.5)
            if r_idx == 0:
                run.font.bold = True
    set_table_width(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def apply_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8 if style_name == "Normal" else 4)
        style.paragraph_format.line_spacing = 1.15
    for level, size, color in [
        (1, 20, RGBColor(0, 0, 0)),
        (2, 16, RGBColor(0, 0, 0)),
        (3, 14, RGBColor(67, 67, 67)),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color


def build():
    doc = Document()
    apply_styles(doc)
    text = SOURCE.read_text(encoding="utf-8")
    lines = []
    in_frontmatter = False
    for raw in text.splitlines():
        if raw.strip() == "---":
            in_frontmatter = not in_frontmatter
            continue
        if in_frontmatter:
            continue
        lines.append(raw.rstrip())

    title_added = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(clean_inline(line[2:]))
            run.font.name = "Arial"
            run.font.size = Pt(26)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            title_added = True
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("- [ ] "):
            add_bullet(doc, "[ ] " + line[6:])
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif re.match(r"^\d+\. ", line):
            items, new_i = parse_numbered_list(lines, i)
            num_id = create_numbering(doc)
            for item in items:
                add_number(doc, item, num_id)
            i = new_i - 1
        elif line.startswith("> "):
            add_para(doc, line[2:], italic=True)
        elif line.startswith("|"):
            rows, new_i = parse_table(lines, i)
            add_table(doc, rows)
            i = new_i - 1
        else:
            add_para(doc, line)
        i += 1

    if not title_added:
        raise RuntimeError("No title found")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
